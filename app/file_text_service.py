import os
import re
import csv
import json
import unicodedata
from html import unescape
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer


MAX_EXTRACTED_CHARS = 50000
DEFAULT_FILE_TITLE_WEIGHT = int(os.getenv("FILE_TITLE_WEIGHT", "4"))
SUPPORTED_FILE_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".docx",
    ".csv",
    ".json",
    ".html",
    ".htm",
    ".rtf",
}

GENERIC_FILE_TITLES = {
    "presentacion de powerpoint",
    "powerpoint presentation",
    "microsoft powerpoint presentation",
    "presentacion",
    "presentation",
    "document",
    "untitled",
    "untitled document",
    "slides",
}

SPANISH_STOPWORDS = {
    "de", "la", "que", "el", "en", "y", "a", "los", "del", "se", "las", "por", "un", "para", "con",
    "no", "una", "su", "al", "lo", "como", "más", "pero", "sus", "le", "ya", "o", "este", "sí", "porque",
    "esta", "entre", "cuando", "muy", "sin", "sobre", "también", "me", "hasta", "hay", "donde", "quien", "desde",
    "todo", "nos", "durante", "todos", "uno", "les", "ni", "contra", "otros", "ese", "eso", "ante", "ellos",
    "e", "esto", "mí", "antes", "algunos", "qué", "unos", "yo", "otro", "otras", "otra", "él", "tanto",
    "esa", "estos", "mucho", "quienes", "nada", "muchos", "cual", "poco", "ella", "estar", "estas", "algunas",
    "algo", "nosotros", "mi", "mis", "tú", "te", "ti", "tu", "tus", "ellas", "nosotras", "vosotros", "vosotras",
    "os", "mío", "mía", "míos", "mías", "tuyo", "tuya", "tuyos", "tuyas", "suyo", "suya", "suyos", "suyas",
    "nuestro", "nuestra", "nuestros", "nuestras", "vuestro", "vuestra", "vuestros", "vuestras", "esos", "esas",
    "estoy", "estás", "está", "estamos", "están", "fue", "fueron", "ser", "es", "son", "era", "eran", "ha",
    "han", "haber", "si", "solo", "sólo", "cada", "además", "través", "puede", "pueden", "hacia", "tras", "aquí",
}


def extract_text_from_file(file_bytes: bytes, extension: str) -> str:
    normalized_extension = extension.lower()

    if normalized_extension == ".pdf":
        return _extract_pdf_text(file_bytes)

    if normalized_extension == ".docx":
        return _extract_docx_text(file_bytes)

    if normalized_extension in {".csv"}:
        return _extract_csv_text(file_bytes)

    if normalized_extension in {".json"}:
        return _extract_json_text(file_bytes)

    if normalized_extension in {".html", ".htm"}:
        return _extract_html_text(file_bytes)

    if normalized_extension in {".rtf"}:
        return _extract_rtf_text(file_bytes)

    if normalized_extension in {".txt", ".md", ".markdown"}:
        text = file_bytes.decode("utf-8", errors="ignore")
        return _clean_limit(text)

    raise ValueError("Formato de archivo no soportado")


def extract_title_from_file(file_bytes: bytes, extension: str, filename: str) -> str:
    normalized_extension = extension.lower()
    filename_title = _title_from_filename(filename)

    if normalized_extension == ".pdf":
        try:
            reader = PdfReader(BytesIO(file_bytes))
            metadata_title = _clean_title(str(getattr(reader.metadata, "title", "") or ""))
            if metadata_title and not _is_generic_title(metadata_title):
                return metadata_title

            for page in reader.pages[:2]:
                page_text = page.extract_text() or ""
                for line in page_text.splitlines():
                    maybe_title = _clean_title(line)
                    if len(maybe_title.split()) >= 4 and not _is_generic_title(maybe_title):
                        return maybe_title
        except Exception:
            return filename_title

    if normalized_extension == ".docx":
        try:
            doc = Document(BytesIO(file_bytes))
            core_title = _clean_title(str(getattr(doc.core_properties, "title", "") or ""))
            if core_title:
                return core_title
            for paragraph in doc.paragraphs[:20]:
                maybe_title = _clean_title(paragraph.text)
                if len(maybe_title.split()) >= 3:
                    return maybe_title
        except Exception:
            return filename_title

    if normalized_extension in {".txt", ".md", ".markdown"}:
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            for line in text.splitlines()[:20]:
                maybe_title = _clean_title(line)
                if len(maybe_title.split()) >= 3:
                    return maybe_title
        except Exception:
            return filename_title

    return filename_title


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages_text: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(page_text)

    return _clean_limit("\n".join(pages_text))


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text and paragraph.text.strip()]
    return _clean_limit("\n".join(paragraphs))


def _extract_csv_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(decoded.splitlines())
    rows: list[str] = []
    for row in reader:
        clean_cells = [cell.strip() for cell in row if cell and cell.strip()]
        if clean_cells:
            rows.append(" | ".join(clean_cells))
    return _clean_limit("\n".join(rows))


def _extract_json_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    try:
        parsed = json.loads(decoded)
    except Exception:
        return _clean_limit(decoded)

    flattened: list[str] = []

    def walk(value) -> None:
        if isinstance(value, dict):
            for key, nested in value.items():
                flattened.append(str(key))
                walk(nested)
            return
        if isinstance(value, list):
            for nested in value:
                walk(nested)
            return
        if value is None:
            return
        flattened.append(str(value))

    walk(parsed)
    return _clean_limit(" ".join(flattened))


def _extract_html_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    no_script = re.sub(r"<script.*?>.*?</script>", " ", decoded, flags=re.IGNORECASE | re.DOTALL)
    no_style = re.sub(r"<style.*?>.*?</style>", " ", no_script, flags=re.IGNORECASE | re.DOTALL)
    no_tags = re.sub(r"<[^>]+>", " ", no_style)
    return _clean_limit(unescape(no_tags))


def _extract_rtf_text(file_bytes: bytes) -> str:
    decoded = file_bytes.decode("utf-8", errors="ignore")
    without_controls = re.sub(r"\\[a-zA-Z]+\d* ?", " ", decoded)
    without_braces = re.sub(r"[{}]", " ", without_controls)
    return _clean_limit(unescape(without_braces))


def _clean_limit(text: str) -> str:
    without_nul = text.replace("\x00", " ")
    normalized = " ".join(without_nul.split())
    return normalized[:MAX_EXTRACTED_CHARS].strip()


def build_index_text_for_clustering(
    text: str,
    title: str | None = None,
    title_weight: int = DEFAULT_FILE_TITLE_WEIGHT,
    chunk_words: int = 280,
    max_chunks: int = 8,
) -> str:
    words = text.split()
    if not words:
        return ""

    chunks: list[str] = []
    start = 0
    while start < len(words) and len(chunks) < max_chunks:
        end = min(start + chunk_words, len(words))
        chunks.append(" ".join(words[start:end]))
        start = end

    sampled_chunks: list[str] = []
    for chunk in chunks:
        sampled_chunks.append(" ".join(chunk.split()[:90]))

    body = "\n\n".join(sampled_chunks).strip()
    clean_title = _clean_title(title or "")

    if not clean_title:
        return body

    safe_weight = max(1, min(int(title_weight), 12))
    weighted_title = "\n".join([clean_title] * safe_weight)
    return f"{weighted_title}\n\n{body}".strip()


def _title_from_filename(filename: str) -> str:
    stem = Path(filename).stem
    normalized = stem.replace("_", " ").replace("-", " ")
    return _clean_title(normalized) or "Untitled Document"


def _clean_title(value: str) -> str:
    return _clean_limit(value)[:220]


def _is_generic_title(title: str) -> bool:
    normalized = unicodedata.normalize("NFKD", title.lower())
    ascii_only = normalized.encode("ascii", "ignore").decode("ascii")
    compact = re.sub(r"[^a-z0-9\s]", " ", ascii_only)
    compact = " ".join(compact.split())
    return compact in GENERIC_FILE_TITLES


def build_semantic_views_for_file(text: str, title: str) -> dict[str, object]:
    clean_title = _clean_title(title)
    summary = _build_file_summary(text)
    keywords = _extract_keywords_tfidf(text, max_keywords=8)
    if not keywords:
        keywords = _extract_keywords_frequency(text, max_keywords=8)
    domain = _infer_domain_from_keywords(keywords)
    topic = clean_title or (keywords[0] if keywords else "documento")
    intent = "referencia documental"

    expanded_parts = [topic]
    if domain:
        expanded_parts.append(domain)
    if summary:
        expanded_parts.append(summary)
    if keywords:
        expanded_parts.append("palabras clave: " + ", ".join(keywords[:6]))

    expanded_context = _clean_limit(". ".join(part for part in expanded_parts if part))[:400]

    return {
        "topic": topic,
        "domain": domain,
        "summary": summary,
        "keywords": keywords,
        "intent": intent,
        "expanded_context": expanded_context,
    }


def build_file_embedding_seed(semantic_views: dict[str, object], title: str) -> str:
    topic = _clean_limit(str(semantic_views.get("topic") or title or "documento"))[:180]
    domain = _clean_limit(str(semantic_views.get("domain") or "documentación general"))[:120]
    summary = _clean_limit(str(semantic_views.get("summary") or ""))[:380]
    intent = _clean_limit(str(semantic_views.get("intent") or "referencia documental"))[:90]
    keywords = semantic_views.get("keywords") if isinstance(semantic_views.get("keywords"), list) else []
    clean_keywords = [_clean_limit(str(keyword))[:40] for keyword in keywords if str(keyword).strip()][:10]

    parts = [
        f"tema: {topic}",
        f"dominio: {domain}",
        f"resumen: {summary}" if summary else "",
        f"intencion: {intent}",
        f"keywords: {', '.join(clean_keywords)}" if clean_keywords else "",
    ]
    return _clean_limit("\n".join(part for part in parts if part))


def _build_file_summary(text: str, max_sentences: int = 3, max_chars: int = 380) -> str:
    if not text:
        return ""

    sentences = [segment.strip() for segment in re.split(r"(?<=[.!?])\s+", text) if segment.strip()]
    selected: list[str] = []
    current_chars = 0

    for sentence in sentences:
        normalized = _clean_limit(sentence)
        if len(normalized) < 25:
            continue
        proposed = current_chars + len(normalized) + (1 if selected else 0)
        if proposed > max_chars:
            break
        selected.append(normalized)
        current_chars = proposed
        if len(selected) >= max_sentences:
            break

    if not selected:
        return _clean_limit(text)[:max_chars]
    return " ".join(selected)


def _extract_keywords_frequency(text: str, max_keywords: int = 8) -> list[str]:
    tokens = re.findall(r"[a-záéíóúñü]{4,}", text.lower())
    frequencies: dict[str, int] = {}
    for token in tokens:
        if token in SPANISH_STOPWORDS:
            continue
        frequencies[token] = frequencies.get(token, 0) + 1

    sorted_tokens = sorted(frequencies.items(), key=lambda item: (-item[1], item[0]))
    return [token for token, _ in sorted_tokens[:max_keywords]]


def _extract_keywords_tfidf(text: str, max_keywords: int = 8) -> list[str]:
    cleaned = _clean_limit(text)
    if not cleaned:
        return []

    chunks = _split_for_tfidf(cleaned, chunk_words=140, max_chunks=24)
    if not chunks:
        return []

    try:
        vectorizer = TfidfVectorizer(
            stop_words=list(SPANISH_STOPWORDS),
            token_pattern=r"(?u)\b[a-záéíóúñü]{4,}\b",
            ngram_range=(1, 2),
            max_features=1200,
        )
        matrix = vectorizer.fit_transform(chunks)
    except Exception:
        return []

    if matrix.shape[1] == 0:
        return []

    scores = matrix.mean(axis=0).A1
    terms = vectorizer.get_feature_names_out()
    top_indices = scores.argsort()[-max_keywords * 3 :][::-1]

    selected: list[str] = []
    for idx in top_indices:
        term = terms[idx].strip()
        if not term:
            continue
        if term in SPANISH_STOPWORDS:
            continue
        if term in selected:
            continue
        selected.append(term)
        if len(selected) >= max_keywords:
            break
    return selected


def _split_for_tfidf(text: str, chunk_words: int, max_chunks: int) -> list[str]:
    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(words) and len(chunks) < max_chunks:
        end = min(start + chunk_words, len(words))
        chunk = " ".join(words[start:end]).strip()
        if chunk:
            chunks.append(chunk)
        start = end
    return chunks


def _infer_domain_from_keywords(keywords: list[str]) -> str:
    joined = " ".join(keywords)
    if any(term in joined for term in {"software", "código", "programación", "algoritmo", "datos", "sistema"}):
        return "tecnología y software"
    if any(term in joined for term in {"empresa", "mercado", "cliente", "finanzas", "negocio", "estrategia"}):
        return "negocio y estrategia"
    if any(term in joined for term in {"salud", "clínico", "médico", "paciente", "tratamiento"}):
        return "salud y medicina"
    if any(term in joined for term in {"investigación", "estudio", "metodología", "análisis", "científico"}):
        return "investigación"
    return "documentación general"
